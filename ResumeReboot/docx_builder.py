from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    set_font(run, size=11, bold=True, color=(31, 73, 125))
    # Add rule as bottom border on same paragraph instead of separate paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_resume_docx(data: dict, job_title: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    r = data.get('rewritten', {})

    contact = r.get('contact', '')
    if contact:
        lines = [l.strip() for l in contact.split('\n') if l.strip()]
        if lines:
            name_p = doc.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_p.paragraph_format.space_after = Pt(2)
            name_run = name_p.add_run(lines[0])
            set_font(name_run, size=18, bold=True, color=(31, 73, 125))
            if len(lines) > 1:
                contact_p = doc.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_p.paragraph_format.space_after = Pt(6)
                contact_run = contact_p.add_run(' · '.join(lines[1:]))
                set_font(contact_run, size=10, color=(80, 80, 80))

    summary = r.get('summary', '')
    if summary:
        add_section_heading(doc, 'Professional Summary')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(summary)
        set_font(run, size=10.5)

    experience = r.get('experience', [])
    if experience:
        add_section_heading(doc, 'Experience')
        for exp in experience:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(8)
            title_p.paragraph_format.space_after = Pt(1)
            title_run = title_p.add_run(exp.get('title', ''))
            set_font(title_run, size=11, bold=True)

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_before = Pt(0)
            meta_p.paragraph_format.space_after = Pt(4)
            company = exp.get('company', '')
            dates = exp.get('dates', '')
            meta_text = f"{company}  |  {dates}" if company and dates else company or dates
            meta_run = meta_p.add_run(meta_text)
            set_font(meta_run, size=10, color=(100, 100, 100))

            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)
                run = bp.add_run(bullet)
                set_font(run, size=10.5)

    education = r.get('education', '')
    if education and 'not provided' not in education.lower() and len(education.strip()) > 10:
        add_section_heading(doc, 'Education')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(education)
        set_font(run, size=10.5)

    skills = r.get('skills', [])
    if skills:
        add_section_heading(doc, 'Technical Skills')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(' · '.join(skills))
        set_font(run, size=10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
